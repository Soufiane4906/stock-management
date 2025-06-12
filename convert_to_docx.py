#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pour convertir le rapport académique Markdown en DOCX
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_academic_report():
    """Crée le rapport académique en DOCX"""
    
    # Créer un nouveau document Word
    doc = Document()
    
    # Titre principal
    title = doc.add_heading('RAPPORT ACADÉMIQUE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre
    subtitle = doc.add_heading('Système de Gestion de Stock Stc-MNGM v2.0', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informations de l'étudiant
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run('Étudiant : ').bold = True
    info_para.add_run('Wissal Benkirane')
    info_para.add_run('\nInstitution : ').bold = True
    info_para.add_run('[Nom de votre institution]')
    info_para.add_run('\nDépartement : ').bold = True
    info_para.add_run('Informatique / Génie Logiciel')
    info_para.add_run('\nDate : ').bold = True
    info_para.add_run('[Date actuelle]')
    info_para.add_run('\nEncadrant : ').bold = True
    info_para.add_run('[Nom de votre encadrant]')
    
    # Table des matières
    doc.add_heading('TABLE DES MATIÈRES', level=1)
    toc = [
        '1. Introduction',
        '2. Contexte et Problématique',
        '3. Objectifs du Projet',
        '4. Analyse des Besoins',
        '5. Conception du Système',
        '6. Architecture Technique',
        '7. Implémentation',
        '8. Tests et Validation',
        '9. Résultats et Discussion',
        '10. Conclusion et Perspectives',
        '11. Références',
        '12. Annexes'
    ]
    
    for item in toc:
        doc.add_paragraph(item, style='List Number')
    
    # 1. INTRODUCTION
    doc.add_heading('1. INTRODUCTION', level=1)
    
    doc.add_heading('1.1 Présentation du Projet', level=2)
    p = doc.add_paragraph()
    p.add_run('Ce rapport présente le développement d\'un système de gestion de stock moderne et sécurisé, dénommé ').bold = False
    p.add_run('Stc-MNGM v2.0').bold = True
    p.add_run('. Ce projet s\'inscrit dans le cadre de l\'étude des systèmes d\'information et de la gestion d\'entreprise, visant à automatiser et optimiser les processus de gestion des stocks, des ventes et des commandes.')
    
    doc.add_heading('1.2 Motivation', level=2)
    p = doc.add_paragraph()
    p.add_run('La gestion manuelle des stocks présente de nombreux défis pour les entreprises :')
    
    challenges = [
        'Risque d\'erreurs humaines',
        'Difficulté de suivi en temps réel',
        'Manque de traçabilité des transactions',
        'Complexité de la génération de rapports',
        'Problèmes de sécurité et d\'accès'
    ]
    
    for challenge in challenges:
        doc.add_paragraph('• ' + challenge, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('L\'objectif de ce projet est de proposer une solution informatique robuste et intuitive pour répondre à ces problématiques.')
    
    # 2. CONTEXTE ET PROBLÉMATIQUE
    doc.add_heading('2. CONTEXTE ET PROBLÉMATIQUE', level=1)
    
    doc.add_heading('2.1 Contexte d\'Étude', level=2)
    p = doc.add_paragraph()
    p.add_run('Dans un environnement économique en constante évolution, les entreprises doivent optimiser leurs processus de gestion pour maintenir leur compétitivité. La gestion des stocks constitue un élément crucial de la chaîne logistique, impactant directement les coûts opérationnels et la satisfaction client.')
    
    doc.add_heading('2.2 Problématique Identifiée', level=2)
    p = doc.add_paragraph()
    p.add_run('Question principale : ').bold = True
    p.add_run('Comment concevoir et implémenter un système de gestion de stock sécurisé, intuitif et performant qui répond aux besoins modernes des entreprises ?')
    
    p = doc.add_paragraph()
    p.add_run('Sous-questions :')
    sub_questions = [
        'Comment assurer la sécurité des données et l\'authentification des utilisateurs ?',
        'Comment optimiser l\'interface utilisateur pour une adoption rapide ?',
        'Comment garantir la cohérence et l\'intégrité des données ?',
        'Comment fournir des outils d\'analyse et de reporting efficaces ?'
    ]
    
    for question in sub_questions:
        doc.add_paragraph('• ' + question, style='List Bullet')
    
    # 3. OBJECTIFS DU PROJET
    doc.add_heading('3. OBJECTIFS DU PROJET', level=1)
    
    doc.add_heading('3.1 Objectif Principal', level=2)
    p = doc.add_paragraph()
    p.add_run('Développer un système de gestion de stock complet et professionnel permettant l\'automatisation des processus de gestion des articles, des ventes, des commandes et des relations avec les clients et fournisseurs.')
    
    doc.add_heading('3.2 Objectifs Spécifiques', level=2)
    
    objectives = [
        ('Sécurité et Authentification', [
            'Implémenter un système d\'authentification sécurisé',
            'Gérer les rôles et permissions utilisateurs',
            'Protéger contre les injections SQL et attaques XSS'
        ]),
        ('Gestion des Articles', [
            'CRUD complet pour les articles',
            'Gestion des catégories',
            'Suivi des quantités en stock',
            'Gestion des dates de fabrication et d\'expiration'
        ]),
        ('Gestion des Ventes', [
            'Enregistrement des transactions',
            'Génération de reçus',
            'Suivi des ventes par client',
            'Annulation de ventes'
        ]),
        ('Gestion des Commandes', [
            'Création de commandes fournisseurs',
            'Mise à jour automatique des stocks',
            'Suivi des commandes'
        ]),
        ('Interface Utilisateur', [
            'Design moderne et responsive',
            'Navigation intuitive',
            'Tableaux de bord informatifs',
            'Expérience utilisateur optimisée'
        ])
    ]
    
    for i, (title, items) in enumerate(objectives, 1):
        doc.add_heading(f'{i}. {title}', level=3)
        for item in items:
            doc.add_paragraph('• ' + item, style='List Bullet')
    
    # 4. ANALYSE DES BESOINS
    doc.add_heading('4. ANALYSE DES BESOINS', level=1)
    
    doc.add_heading('4.1 Analyse Fonctionnelle', level=2)
    
    doc.add_heading('4.1.1 Acteurs du Système', level=3)
    actors = [
        ('Administrateur', 'Accès complet à toutes les fonctionnalités'),
        ('Utilisateur', 'Accès limité aux opérations de base'),
        ('Système', 'Gestion automatique des sessions et de la sécurité')
    ]
    
    for actor, description in actors:
        p = doc.add_paragraph()
        p.add_run(f'• {actor} : ').bold = True
        p.add_run(description)
    
    doc.add_heading('4.1.2 Fonctionnalités Principales', level=3)
    
    features = [
        ('Gestion des Utilisateurs', [
            'Connexion/Déconnexion sécurisée',
            'Gestion des profils utilisateurs',
            'Contrôle d\'accès basé sur les rôles'
        ]),
        ('Gestion des Articles', [
            'Ajout, modification, suppression d\'articles',
            'Catégorisation des articles',
            'Gestion des images',
            'Suivi des stocks'
        ]),
        ('Gestion des Ventes', [
            'Création de ventes',
            'Sélection d\'articles et de clients',
            'Calcul automatique des prix',
            'Génération de reçus'
        ]),
        ('Gestion des Commandes', [
            'Création de commandes fournisseurs',
            'Mise à jour automatique des stocks',
            'Suivi des commandes'
        ]),
        ('Gestion des Relations', [
            'Gestion des clients',
            'Gestion des fournisseurs',
            'Gestion des contacts'
        ]),
        ('Tableaux de Bord', [
            'Statistiques en temps réel',
            'Ventes récentes',
            'Articles les plus vendus',
            'Indicateurs de performance'
        ])
    ]
    
    for feature, items in features:
        p = doc.add_paragraph()
        p.add_run(feature).bold = True
        for item in items:
            doc.add_paragraph('• ' + item, style='List Bullet')
    
    # 5. CONCEPTION DU SYSTÈME
    doc.add_heading('5. CONCEPTION DU SYSTÈME', level=1)
    
    doc.add_heading('5.1 Modèle de Données', level=2)
    
    doc.add_heading('5.1.1 Diagramme Entité-Relation', level=3)
    p = doc.add_paragraph()
    p.add_run('Le système repose sur 8 entités principales :')
    
    entities = [
        ('Users', 'Gestion des utilisateurs et authentification'),
        ('Article', 'Stockage des informations sur les produits'),
        ('CategorieArticle', 'Classification des articles'),
        ('Client', 'Informations sur les clients'),
        ('Fournisseur', 'Informations sur les fournisseurs'),
        ('Vente', 'Enregistrement des transactions de vente'),
        ('Commande', 'Gestion des commandes fournisseurs'),
        ('Contact', 'Gestion des contacts')
    ]
    
    for entity, description in entities:
        p = doc.add_paragraph()
        p.add_run(f'• {entity} : ').bold = True
        p.add_run(description)
    
    # 6. ARCHITECTURE TECHNIQUE
    doc.add_heading('6. ARCHITECTURE TECHNIQUE', level=1)
    
    doc.add_heading('6.1 Technologies Utilisées', level=2)
    
    doc.add_heading('6.1.1 Backend', level=3)
    backend_tech = [
        ('PHP 7.4+', 'Langage de programmation côté serveur'),
        ('MySQL/MariaDB', 'Système de gestion de base de données'),
        ('PDO', 'Interface d\'accès aux données'),
        ('Sessions PHP', 'Gestion des sessions utilisateur')
    ]
    
    for tech, description in backend_tech:
        p = doc.add_paragraph()
        p.add_run(f'• {tech} : ').bold = True
        p.add_run(description)
    
    doc.add_heading('6.1.2 Frontend', level=3)
    frontend_tech = [
        ('HTML5', 'Structure des pages'),
        ('CSS3', 'Styles et mise en forme'),
        ('JavaScript', 'Interactivité côté client'),
        ('Boxicons', 'Bibliothèque d\'icônes')
    ]
    
    for tech, description in frontend_tech:
        p = doc.add_paragraph()
        p.add_run(f'• {tech} : ').bold = True
        p.add_run(description)
    
    # 7. IMPLÉMENTATION
    doc.add_heading('7. IMPLÉMENTATION', level=1)
    
    doc.add_heading('7.1 Configuration de l\'Environnement', level=2)
    
    doc.add_heading('7.1.1 Installation des Prérequis', level=3)
    p = doc.add_paragraph()
    p.add_run('• Installation de XAMPP')
    p.add_run('\n• Configuration de PHP et MySQL')
    p.add_run('\n• Création de la base de données')
    
    # 8. TESTS ET VALIDATION
    doc.add_heading('8. TESTS ET VALIDATION', level=1)
    
    doc.add_heading('8.1 Tests Fonctionnels', level=2)
    
    doc.add_heading('8.1.1 Tests d\'Authentification', level=3)
    auth_tests = [
        'Connexion avec identifiants valides',
        'Rejet des identifiants invalides',
        'Gestion des sessions',
        'Déconnexion sécurisée'
    ]
    
    for test in auth_tests:
        p = doc.add_paragraph()
        p.add_run('✅ ').bold = True
        p.add_run(test)
    
    # 9. RÉSULTATS ET DISCUSSION
    doc.add_heading('9. RÉSULTATS ET DISCUSSION', level=1)
    
    doc.add_heading('9.1 Fonctionnalités Réalisées', level=2)
    
    doc.add_heading('9.1.1 Fonctionnalités Principales', level=3)
    features_realized = [
        'Système d\'authentification complet',
        'Gestion complète des articles',
        'Gestion des ventes et commandes',
        'Interface utilisateur moderne',
        'Tableaux de bord informatifs',
        'Gestion des clients et fournisseurs'
    ]
    
    for feature in features_realized:
        p = doc.add_paragraph()
        p.add_run('✅ ').bold = True
        p.add_run(feature)
    
    # 10. CONCLUSION ET PERSPECTIVES
    doc.add_heading('10. CONCLUSION ET PERSPECTIVES', level=1)
    
    doc.add_heading('10.1 Bilan du Projet', level=2)
    p = doc.add_paragraph()
    p.add_run('Ce projet de système de gestion de stock a permis de développer une solution complète et professionnelle répondant aux besoins modernes des entreprises. Les objectifs fixés ont été atteints avec succès :')
    
    achievements = [
        'Système d\'authentification sécurisé',
        'Gestion complète des stocks',
        'Interface utilisateur moderne',
        'Performance optimisée',
        'Sécurité renforcée'
    ]
    
    for achievement in achievements:
        p = doc.add_paragraph()
        p.add_run('✅ ').bold = True
        p.add_run(achievement)
    
    # 11. RÉFÉRENCES
    doc.add_heading('11. RÉFÉRENCES', level=1)
    
    doc.add_heading('11.1 Documentation Technique', level=2)
    refs = [
        'Documentation PHP : https://www.php.net/docs.php',
        'Documentation MySQL : https://dev.mysql.com/doc/',
        'Documentation PDO : https://www.php.net/manual/fr/book.pdo.php',
        'Guide CSS Grid : https://developer.mozilla.org/fr/docs/Web/CSS/CSS_Grid_Layout'
    ]
    
    for ref in refs:
        doc.add_paragraph(ref, style='List Bullet')
    
    # 12. ANNEXES
    doc.add_heading('12. ANNEXES', level=1)
    
    annexes = [
        'Annexe A : Schéma de Base de Données',
        'Annexe B : Captures d\'Écran',
        'Annexe C : Code Source',
        'Annexe D : Tests'
    ]
    
    for annexe in annexes:
        doc.add_paragraph(annexe, style='List Bullet')
    
    # Signature
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Rapport rédigé par : ').bold = True
    p.add_run('Wissal Benkirane')
    p.add_run('\nDate de finalisation : ').bold = True
    p.add_run('[Date]')
    p.add_run('\nVersion : ').bold = True
    p.add_run('1.0')
    
    # Sauvegarder le document
    doc.save('Rapport_Academique_Stc-MNGM_v2.0.docx')
    print("Document créé avec succès : Rapport_Academique_Stc-MNGM_v2.0.docx")

if __name__ == "__main__":
    create_academic_report() 